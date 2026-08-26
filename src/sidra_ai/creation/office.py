"""Write a generated deck out as a real Word, Excel or PowerPoint file.

The HTML deck is the artifact that always exists. These formats are optional
in the strict sense: the packages live in the ``creation`` extra, an operator
may not have installed them, and every writer here returns *whether it wrote
and why not* rather than raising or - worse - reporting success for a file it
did not produce. A deck path that claimed three formats and delivered one is
the failure this module is shaped around.

Nothing new is invented in the conversion. The content is the deck the
grounding pipeline already built, blanks included: a number missing from the
corpus stays 〔社長が埋める欄〕 in Word exactly as it does in HTML. A format
change is not an occasion to fill anything in.

What the checker proves, and what it does not: ``validate_office`` opens the
file as the ZIP package it is, requires the parts the format is defined by,
and parses every XML member. That catches a truncated write, a missing
content-type map, and malformed markup. It does **not** prove Microsoft Word
opens it - nothing in this container can - so the number it feeds is named
for structure, and this paragraph is the reason it is not called "compatible".
"""

from __future__ import annotations

import zipfile
from pathlib import Path
from xml.etree import ElementTree

from sidra_ai.creation.decks import GeneratedDeck, save_pptx

#: Words that name a format. Literal, like every other matcher in this
#: package: a loose one finds a format request in every sentence.
FORMAT_WORDS: dict[str, tuple[str, ...]] = {
    "docx": ("word", "ワード", "docx", "文書ファイル"),
    "xlsx": ("excel", "エクセル", "xlsx", "表計算", "スプレッドシート"),
    "pptx": ("pptx", "powerpoint", "パワポ", "パワーポイント", "スライド"),
}

#: The members every valid package of each format must contain.
REQUIRED_PARTS: dict[str, tuple[str, ...]] = {
    "docx": ("[Content_Types].xml", "word/document.xml"),
    "xlsx": ("[Content_Types].xml", "xl/workbook.xml"),
    "pptx": ("[Content_Types].xml", "ppt/presentation.xml"),
}


def requested_formats(request: str) -> tuple[str, ...]:
    """Which Office formats the wording asked for, in a stable order."""

    lowered = request.casefold()
    return tuple(
        fmt
        for fmt in ("docx", "xlsx", "pptx")
        if any(word in lowered for word in FORMAT_WORDS[fmt])
    )


def _rows(deck: GeneratedDeck) -> list[tuple[str, str, str]]:
    """Flatten the deck to (slide, bullet, sources) once, for every writer.

    One flattening rather than three keeps the formats from disagreeing about
    what the deck said - which is the quiet way a conversion starts editing.
    """

    rows: list[tuple[str, str, str]] = []
    for slide in deck.slides:
        if not slide.bullets:
            rows.append((slide.title, "", "; ".join(slide.sources)))
            continue
        for bullet in slide.bullets:
            rows.append((slide.title, bullet, "; ".join(slide.sources)))
    return rows


def write_docx(deck: GeneratedDeck, path: str | Path) -> tuple[bool, str]:
    try:
        from docx import Document  # type: ignore[import-not-found]
    except Exception as exc:  # noqa: BLE001 - absence is an expected case
        return False, f"python-docx not available: {type(exc).__name__}"

    document = Document()
    document.add_heading(deck.title, level=0)
    for slide in deck.slides:
        document.add_heading(slide.title, level=1)
        for bullet in slide.bullets:
            document.add_paragraph(bullet, style="List Bullet")
        if slide.sources:
            document.add_paragraph("出典: " + "; ".join(slide.sources))
    if deck.unfilled:
        document.add_heading("空欄のまま残した節", level=1)
        for title in deck.unfilled:
            document.add_paragraph(title, style="List Bullet")
    document.save(str(path))
    return True, "python-docx"


def write_xlsx(deck: GeneratedDeck, path: str | Path) -> tuple[bool, str]:
    try:
        from openpyxl import Workbook  # type: ignore[import-not-found]
    except Exception as exc:  # noqa: BLE001 - absence is an expected case
        return False, f"openpyxl not available: {type(exc).__name__}"

    book = Workbook()
    sheet = book.active
    sheet.title = "deck"
    # The deck's name, first: a sheet of rows with no title is a table that
    # has lost which document it came from, and the other two formats carry
    # it. Formats of one deck disagreeing about its name is the drift this
    # module's single flattening exists to prevent.
    sheet.append((deck.title,))
    sheet.append(("節", "項目", "出典"))
    for row in _rows(deck):
        sheet.append(row)
    book.save(str(path))
    return True, "openpyxl"


def write_office(deck: GeneratedDeck, directory: str | Path, stem: str) -> dict[str, dict]:
    """Write every format, and report each one's outcome separately.

    A single boolean for "the Office export worked" would hide the case this
    is most likely to hit: two packages installed and one not.
    """

    target = Path(directory)
    target.mkdir(parents=True, exist_ok=True)
    writers = {
        "docx": write_docx,
        "xlsx": write_xlsx,
        "pptx": save_pptx,
    }
    results: dict[str, dict] = {}
    for fmt, writer in writers.items():
        path = target / f"{stem}.{fmt}"
        written, why = writer(deck, path)
        results[fmt] = {
            "written": written,
            "reason": why,
            "path": str(path) if written else "",
            "valid": validate_office(path, fmt)["valid"] if written else False,
        }
    return results


def validate_office(path: str | Path, fmt: str) -> dict:
    """Open the package and prove it is one, part by part.

    Structural only - see the module docstring. A file that passes here is a
    well-formed OOXML package with the members its format is defined by; it
    is not a claim about any particular application.
    """

    failures: list[str] = []
    target = Path(path)
    if not target.is_file():
        return {"valid": False, "failures": [f"{target.name}: not written"]}
    try:
        with zipfile.ZipFile(target) as package:
            names = set(package.namelist())
            for part in REQUIRED_PARTS.get(fmt, ()):
                if part not in names:
                    failures.append(f"missing {part}")
            for name in names:
                if not name.endswith(".xml") and not name.endswith(".rels"):
                    continue
                try:
                    ElementTree.fromstring(package.read(name))
                except ElementTree.ParseError as exc:
                    failures.append(f"{name}: {exc}")
    except zipfile.BadZipFile as exc:
        failures.append(f"not a package: {exc}")
    return {"valid": not failures, "failures": failures}


__all__ = [
    "FORMAT_WORDS",
    "REQUIRED_PARTS",
    "requested_formats",
    "validate_office",
    "write_docx",
    "write_office",
    "write_xlsx",
]
